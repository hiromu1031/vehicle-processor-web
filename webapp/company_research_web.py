"""
企業情報Web調査プロセッサ（Web版）
会社名・住所・代表者名から企業情報をWeb検索して収集し、Word文書にまとめる
"""
from io import BytesIO
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from utils.claude_client import ClaudeClient


class CompanyResearchWeb:
    """企業情報Web調査プロセッサ（Web版）"""

    def __init__(self):
        self.claude_client = ClaudeClient()

    def search_company_info(
        self,
        company_name: str,
        address: str,
        representative: str
    ) -> Dict[str, Any]:
        """
        Web検索で企業情報を収集

        Args:
            company_name: 会社名
            address: 所在地
            representative: 代表者名

        Returns:
            収集した情報の辞書
        """
        # 検索クエリのリスト
        search_queries = [
            f"{company_name} 会社概要",
            f"{company_name} 事業内容",
            f"{company_name} 採用情報 求人",
            f"{company_name} {representative} インタビュー",
            f"{company_name} {representative} 社長",
            f"{company_name} ニュース",
            f"{company_name} プレスリリース",
        ]

        # 各クエリで検索を実行
        search_results = {}
        for query in search_queries:
            try:
                # WebSearchを実行
                result = self._web_search(query)
                search_results[query] = result
            except Exception as e:
                print(f"検索エラー ({query}): {str(e)}")
                search_results[query] = {"error": str(e)}

        # 情報を構造化
        structured_info = self._structure_information(
            company_name,
            address,
            representative,
            search_results
        )

        return structured_info

    def _web_search(self, query: str) -> Dict[str, Any]:
        """
        Claudeのトレーニングデータから企業情報を抽出（簡易版）

        Note: リアルタイムのWeb検索ではなく、2025年1月までのトレーニングデータから情報を抽出
        """
        prompt = f"""
以下の検索クエリについて、あなたのトレーニングデータ（2025年1月まで）に含まれる情報を教えてください。

検索クエリ: {query}

以下の形式で情報を提供してください：
1. 主要な情報の要約（具体的な事実のみ、推測は含めない）
2. 情報がない場合は「情報なし」と明記

JSONフォーマット:
{{
  "summary": "情報の要約（情報がない場合は空文字列）",
  "has_info": true/false,
  "confidence": "high/medium/low"
}}

重要:
- 確実な情報のみを含める
- 推測や一般論は含めない
- 情報がない場合は正直に「情報なし」とする
"""

        try:
            # Claude APIで情報を取得
            result = self.claude_client.client.messages.create(
                model=self.claude_client.model,
                max_tokens=1500,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )

            response_text = result.content[0].text

            # JSONをパース
            import json
            if "```json" in response_text:
                json_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                json_text = response_text.split("```")[1].split("```")[0].strip()
            else:
                json_text = response_text.strip()

            parsed = json.loads(json_text)

            # 出典情報を追加（簡易版）
            return {
                "summary": parsed.get("summary", ""),
                "sources": [
                    {
                        "url": "N/A",
                        "title": "Claude AI トレーニングデータ",
                        "type": "AIトレーニングデータ（2025年1月まで）"
                    }
                ] if parsed.get("has_info") else [],
                "has_info": parsed.get("has_info", False)
            }

        except Exception as e:
            print(f"情報抽出エラー ({query}): {str(e)}")
            return {
                "summary": "",
                "sources": [],
                "has_info": False,
                "error": str(e)
            }

    def _structure_information(
        self,
        company_name: str,
        address: str,
        representative: str,
        search_results: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        検索結果を構造化された情報に変換
        """
        structured = {
            "company_name": company_name,
            "address": address,
            "representative": representative,
            "sections": []
        }

        # セクションごとに情報を整理
        section_map = {
            "会社概要・事業内容": ["会社概要", "事業内容"],
            "採用情報": ["採用情報"],
            "代表者インタビュー": ["インタビュー", "社長"],
            "ニュース・プレスリリース": ["ニュース", "プレスリリース"],
        }

        for section_title, keywords in section_map.items():
            section_content = []
            section_sources = []

            for query, result in search_results.items():
                if any(keyword in query for keyword in keywords):
                    if "error" not in result:
                        section_content.append(result.get("summary", ""))
                        section_sources.extend(result.get("sources", []))

            if section_content:
                structured["sections"].append({
                    "title": section_title,
                    "content": "\n\n".join(section_content),
                    "sources": section_sources
                })

        return structured

    def generate_word_report(
        self,
        company_info: Dict[str, Any]
    ) -> BytesIO:
        """
        収集した情報からWord文書を生成

        Args:
            company_info: search_company_info()の結果

        Returns:
            BytesIO: Word文書のバイトストリーム
        """
        doc = Document()

        # ドキュメントのスタイル設定
        style = doc.styles['Normal']
        style.font.name = 'Yu Gothic'
        style.font.size = Pt(11)

        # タイトル
        title = doc.add_heading(f'{company_info["company_name"]} 企業情報調査レポート', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 作成日
        date_para = doc.add_paragraph()
        date_para.add_run(f'作成日: {datetime.now().strftime("%Y年%m月%d日")}').italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()  # 空行

        # 基本情報
        doc.add_heading('■ 基本情報', level=1)
        doc.add_paragraph(f'会社名: {company_info["company_name"]}')
        doc.add_paragraph(f'所在地: {company_info["address"]}')
        doc.add_paragraph(f'代表者: {company_info["representative"]}')

        doc.add_paragraph()  # 空行

        # 各セクション
        for section in company_info["sections"]:
            doc.add_heading(f'■ {section["title"]}', level=1)

            # 内容
            content_para = doc.add_paragraph(section["content"])

            doc.add_paragraph()  # 空行

            # 出典
            if section["sources"]:
                sources_heading = doc.add_paragraph()
                sources_heading.add_run('📎 出典:').bold = True

                for source in section["sources"]:
                    source_para = doc.add_paragraph(style='List Bullet')
                    source_para.add_run(f'{source.get("title", "不明")} ')
                    source_para.add_run(f'({source.get("type", "Web")})').italic = True
                    # URLがN/A以外の場合のみ表示
                    if source.get("url") and source.get("url") != "N/A":
                        source_para.add_run(f'\n   {source.get("url", "")}')

            doc.add_paragraph()  # 空行

        # BytesIOに保存
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        return word_buffer
